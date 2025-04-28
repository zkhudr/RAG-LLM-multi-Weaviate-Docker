#!/usr/bin/env python
import os
import json
import hashlib
import logging
import argparse
import csv
import re
import warnings
import traceback
from pathlib import Path
import datetime as dt
from datetime import timezone
from typing import List, Dict, Optional, Tuple, Union
# For PDF processing, use pdfplumber (as in ingest_docs_v7.py)
import pdfplumber
import PyPDF2
import pandas as pd
from tenacity import retry, stop_after_attempt, wait_exponential
# LangChain components for Document and splitting
from langchain_core.documents import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
# Import configuration from your project and pipeline settings from ingest_docs_v7.py
from config import cfg
from ingest_docs_v7 import PipelineConfig, RobustPDFLoaderV4
############
import hashlib
import sys
from typing import Dict, List, Optional, Any, Generator
# Weaviate v4 imports
import weaviate
from weaviate.classes.config import Property, DataType, Configure,VectorDistances
from weaviate.classes.data import DataObject
from weaviate.exceptions import (
    WeaviateConnectionError,
    WeaviateInvalidInputError,
    WeaviateBatchError,
    WeaviateBaseError
)
# LangChain Components
from langchain_community.vectorstores.utils import filter_complex_metadata
from langchain_ollama import OllamaEmbeddings
# Try to import DOCX support
try:
    import docx
except ImportError:
    docx = None
    logging.warning("python-docx is not installed; DOCX files will not be processed.")

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)
# --- Extraction Functions for Various File Types ---


def extract_text_from_txt(filepath: Path) -> str:
    try:
        with open(filepath, "r", encoding="utf-8") as f:
            return f.read()
    except Exception as e:
        logger.error("Error reading TXT file %s: %s", filepath.name, e)
        return ""

def extract_text_from_docx(filepath: Path) -> str:
    if docx is None:
        logger.error("python-docx is not installed. Cannot process DOCX file: %s", filepath.name)
        return ""
    try:
        doc = docx.Document(filepath)
        paragraphs = [para.text for para in doc.paragraphs]
        return "\n".join(paragraphs)
    except Exception as e:
        logger.error("Error reading DOCX file %s: %s", filepath.name, e)
        return ""

def extract_text_from_csv(filepath: Path) -> str:
    try:
        with open(filepath, "r", encoding="utf-8") as f:
            reader = csv.reader(f)
            rows = list(reader)
            # Join each row with commas and rows with newlines.
            return "\n".join([", ".join(row) for row in rows])
    except Exception as e:
        logger.error("Error reading CSV file %s: %s", filepath.name, e)
        return ""

# Modify extract_text signature and logic
# Add optional client parameter with default None
def extract_text(filepath: Path, client_instance: Optional[weaviate.Client] = None) -> str:
    ext = filepath.suffix.lower()
    if ext == ".txt":
        return extract_text_from_txt(filepath)
    # ... (keep other elif blocks for docx, csv) ...
    elif ext == ".pdf":
        loader_client = None
        try:
            # --- USE PASSED-IN CLIENT if available ---
            if client_instance:
                 loader_client = client_instance
                 logger.debug(f"Using provided client instance for PDF extraction: {filepath.name}")
            else:
                # Fallback: Get a temporary client if none was passed (should be avoided)
                logger.warning(f"No client passed to extract_text for {filepath.name}, getting temporary one.")
                loader_client = PipelineConfig.get_client() # Original problematic path

            if loader_client is None:
                 logger.error(f"Failed to get/use any Weaviate client for PDF extraction: {filepath.name}")
                 return "" # Return empty on client failure

            # Use the determined client (passed-in or temporary)
            loader = RobustPDFLoaderV4(str(filepath), client=loader_client)
            docs = loader.load()
            return "\n".join(doc.page_content for doc in docs if doc.page_content)

        except Exception as e:
            logger.error("Error processing PDF file %s: %s", filepath.name, e)
            return ""
        # --- IMPORTANT: Do NOT close the client here if it was passed in ---
        # The caller (run_ingestion) is responsible for closing the main client.
        # If a temporary client was created, its closure relies on get_client/RobustPDFLoaderV4 internals or GC.

    else:
        logger.warning("Unsupported file type: %s", filepath.name)
        return ""

# Modify load_new_documents signature and call to extract_text
# Add client parameter
class IncrementalDocumentProcessorBlock:
    # ... (keep __init__, load_metadata, save_metadata, compute_hash) ...

    # Add client parameter
    def load_new_documents(self, client_instance: weaviate.Client) -> List[Document]:
        """Iterate over files... Pass client to extract_text."""
        docs = []
        for filepath in self.data_dir.glob("*"):
            # ... (keep file check, hash check logic) ...
            if stored_hash == current_hash:
                continue

            logger.info("Processing new/updated file: %s", filepath.name)
            # --- PASS THE CLIENT ---
            text = extract_text(filepath, client_instance=client_instance)

            if text is None or not text.strip():
                 logger.warning("No text extracted or text is empty/whitespace for %s, skipping.", filepath.name)
                 self.processed_files[str(filepath)] = current_hash
                 continue

            # --- CORRECTED INDENTATION & DOC CREATION ---
            try:
                # ... (logic to create Document object with metadata) ...
                doc = Document(page_content=text, metadata=metadata)
                docs.append(doc)
                self.processed_files[str(filepath)] = current_hash
            except Exception as doc_create_e:
                logger.error(f"Error creating Document object for {filepath.name}: {doc_create_e}")
                continue
            # --- END CORRECTION ---
        return docs

    # Modify process method to pass the client down
    def process(self, client_instance: weaviate.Client) -> List[Document]:
        """Load new documents using the provided client, split, filter, update metadata."""
        # --- PASS THE CLIENT ---
        docs = self.load_new_documents(client_instance=client_instance)
        if not docs:
            logger.info("No new documents to process.")
            return []

        # ... (keep splitter logic) ...
        valid_docs = [d for d in split_docs if len(d.page_content) >= PipelineConfig.MIN_CONTENT_LENGTH]
        logger.info("Prepared %d valid document chunks for insertion.", len(valid_docs))
        self.save_metadata() # Save updated hashes
        return valid_docs

# Modify run_ingestion to pass the main client to the processor
def run_ingestion(folder: str) -> dict:
    client = None
    embeddings_obj = None
    num_processed = 0
    try:
        # 1. Initialize the MAIN Weaviate client
        logger.info("Attempting to get Weaviate client via PipelineConfig...")
        client = PipelineConfig.get_client()

        if client is None:
            logger.critical("Failed to obtain a valid Weaviate client...")
            return { ... } # Error dict

        if not client.is_live():
            logger.critical("Weaviate server is not responsive.")
            return { ... } # Error dict
        logger.info("Weaviate client is live.")

        # 4. Initialize embeddings (as before)
        logger.info("Initializing embeddings...")
        try:
            embeddings_obj = OllamaEmbeddings(...)
            logger.info(...)
        except Exception as embed_e:
            logger.error(...)
            return { ... } # Error dict

        # 5. Process new/modified documents incrementally
        logger.info(f"Starting incremental processing for directory: {folder}")
        incremental_processor = IncrementalDocumentProcessorBlock(data_dir=folder)
        # --- PASS THE MAIN CLIENT to process ---
        new_docs = incremental_processor.process(client_instance=client)
        num_processed = len(new_docs)

        # 6. Insert the resulting document chunks (as before)
        if new_docs:
            logger.info(...)
            store_documents(new_docs, client, embeddings_obj)
        else:
            logger.info(...)

        # 7. Log success and prepare return value (as before)
        logger.info(...)
        return { ... } # Success dict

    except Exception as e:
        logger.critical(...)
        return { ... } # Generic error dict

    finally:
        # 8. Ensure the MAIN client connection is closed (as before)
        if client and hasattr(client, 'close') and callable(client.close):
            try:
                if client.is_connected():
                    client.close()
                    logger.info("Closed Weaviate client connection.")
                # ... (else block) ...
            except Exception as close_err:
                logger.error(...)


def chunk_text(text: str, max_chunk_size: int) -> list:
    """Simple splitting of text into chunks up to max_chunk_size characters.
    For production, you might use a sentence or paragraph splitter instead."""
    chunks = []
    start = 0
    while start < len(text):
        end = min(start + max_chunk_size, len(text))
        chunks.append(text[start:end])
        start = end
    return chunks

# --- Incremental Ingestion Class ---
class IncrementalDocumentProcessorBlock:
    def __init__(self, data_dir: str, meta_filename="ingested_files.json"):
        self.data_dir = Path(data_dir).resolve()
        self.meta_file = self.data_dir / meta_filename
        self.processed_files = self.load_metadata()
    
    def load_metadata(self) -> dict:
        """Load a JSON file mapping file paths to their last processed hash."""
        if self.meta_file.exists():
            try:
                with open(self.meta_file, "r", encoding="utf-8") as f:
                    return json.load(f)
            except Exception as e:
                logger.error("Error loading metadata: %s", e)
                return {}
        return {}
    
    def save_metadata(self) -> None:
        """Save the updated metadata dictionary."""
        try:
            with open(self.meta_file, "w", encoding="utf-8") as f:
                json.dump(self.processed_files, f, indent=2)
        except Exception as e:
            logger.error("Error saving metadata: %s", e)
    
    def compute_hash(self, filepath: Path) -> str:
        """Compute the SHA-256 hash of a file’s contents."""
        hasher = hashlib.sha256()
        try:
            with open(filepath, "rb") as f:
                while True:
                    buf = f.read(65536)
                    if not buf:
                        break
                    hasher.update(buf)
            return hasher.hexdigest()
        except Exception as e:
            logger.error("Error computing hash for %s: %s", filepath.name, e)
            return ""
    
    def load_new_documents(self) -> List[Document]:
        """Iterate over files in the data directory and return Document objects
        for new or updated files only."""
        docs = []
        for filepath in self.data_dir.glob("*"): # Consider rglob("**/*") if subdirs are needed
            if not filepath.is_file():
                continue

            # Check file type support using PipelineConfig class attribute (if defined)
            # Assuming PipelineConfig has FILE_TYPES attribute similar to config.document
            supported_extensions = set(f".{ext.lower()}" for ext in getattr(PipelineConfig, 'FILE_TYPES', ['pdf', 'txt', 'csv', 'docx', 'md']))
            if filepath.suffix.lower() not in supported_extensions:
                # logger.debug(f"Skipping unsupported file type: {filepath.name}") # Make debug?
                continue # Skip unsupported types

            current_hash = self.compute_hash(filepath)
            stored_hash = self.processed_files.get(str(filepath))

            if stored_hash == current_hash:
                # logger.debug("Skipping file (unchanged): %s", filepath.name) # Make debug?
                continue

            logger.info("Processing new/updated file: %s", filepath.name)
            text = extract_text(filepath) # text might be None here

            # --- CORRECTED CHECK ---
            if text is None or not text.strip(): # Explicitly check for None first
                logger.warning("No text extracted or text is empty/whitespace for %s, skipping.", filepath.name)
                # Let's update hash to avoid retrying problematic files indefinitely.
                self.processed_files[str(filepath)] = current_hash
                continue # Skip to the next file

            try:
                # Create a Document object; we use basic metadata
                file_path_obj = Path(filepath) # Ensure it's Path object
                stat = file_path_obj.stat()
                # Use timezone-aware datetime objects and format correctly
                created_dt = dt.datetime.fromtimestamp(stat.st_ctime, tz=timezone.utc).isoformat().replace('+00:00', 'Z')
                modified_dt = dt.datetime.fromtimestamp(stat.st_mtime, tz=timezone.utc).isoformat().replace('+00:00', 'Z')

                metadata = {
                    "source": str(filepath),
                    "filetype": file_path_obj.suffix.lower(),
                    "created_date": created_dt,
                    "modified_date": modified_dt,
                    "page": 0, # Default for whole-doc text, splitter might override
                    # Include other default fields required by schema if any
                    "content_flags": [],
                    "error_messages": []
                }

                doc = Document(page_content=text, metadata=metadata)
                docs.append(doc)

                # Update the metadata record only after successful processing and doc creation
                self.processed_files[str(filepath)] = current_hash

            except Exception as doc_create_e:
                logger.error(f"Error creating Document object for {filepath.name}: {doc_create_e}", exc_info=True)
                # Do NOT update hash if doc creation failed, allow retry later
                continue # Skip to the next file
            # --- END INDENTATION FIX ---
        return docs
    
    def process(self) -> List[Document]:
        """Load new documents, split them into chunks, filter out short chunks, and update metadata."""
        docs = self.load_new_documents()
        if not docs:
            logger.info("No new documents to process.")
            return []
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=PipelineConfig.CHUNK_SIZE,
            chunk_overlap=PipelineConfig.CHUNK_OVERLAP
        )
        split_docs = splitter.split_documents(docs)
        valid_docs = [d for d in split_docs if len(d.page_content) >= PipelineConfig.MIN_CONTENT_LENGTH]
        logger.info("Prepared %d valid document chunks for insertion.", len(valid_docs))
        self.save_metadata()
        return valid_docs
# --- Store Documents in Weaviate ---
def store_documents(docs: List[Document], client: weaviate.Client, embeddings_obj) -> None:
        if not docs:
            logger.info("No documents to insert.")
            return
        collection_name = PipelineConfig.VECTORSTORE_COLLECTION
        try:
            collection = client.collections.get(collection_name)
        except Exception as e:
            logger.error("Error fetching collection '%s': %s", collection_name, e)
            return
        objects = []
        logger.info("Preparing %d document chunks for insertion.", len(docs))

        for i, doc in enumerate(docs):
            try:
                @retry(stop=stop_after_attempt(PipelineConfig.RETRY_ATTEMPTS),
                    wait=wait_exponential(multiplier=PipelineConfig.RETRY_WAIT_MULTIPLIER,
                                            max=PipelineConfig.RETRY_MAX_WAIT))
                def get_embedding():
                    return embeddings_obj.embed_query(doc.page_content)
                embedding = get_embedding()

                file_path = Path(doc.metadata["source"])
                stat = file_path.stat()
                created_dt = dt.datetime.fromtimestamp(stat.st_ctime, tz=timezone.utc).isoformat().replace('+00:00', 'Z')
                modified_dt = dt.datetime.fromtimestamp(stat.st_mtime, tz=timezone.utc).isoformat().replace('+00:00', 'Z')

                properties = {
                    "content": doc.page_content,
                    "source": doc.metadata.get("source", "unknown"),
                    "filetype": doc.metadata.get("filetype", "unknown"),
                    "page": 0,
                    "created_date": None,  # Will fix in next step
                    "modified_date": None, # Will fix in next step
                    "content_flags": [],   # Add this line
                    "error_messages": []   # Add this line
                }
                objects.append(
                    DataObject(
                        properties=properties,
                        vector=embedding
                    )
                )
            except Exception as e:
                logger.error("Error preparing document %d (source: %s): %s", i, doc.metadata.get("source", "N/A"), e)
        if not objects:
            logger.warning("No objects prepared for insertion.")
            return
        try:
            response = collection.data.insert_many(objects=objects)
            if response and response.has_errors:
                logger.error("Insertion completed with errors: %s", response.errors)
            else:
                logger.info("Document chunks inserted successfully.")
        except Exception as e:
            logger.error("Batch insertion failed: %s", e)

def run_ingestion(folder: str) -> dict:
    """
    Run the incremental ingestion process on the specified folder.

    Connects to Weaviate, processes new/updated documents, embeds them,
    and stores them in the Weaviate collection.

    Args:
        folder (str): The path to the directory containing documents.

    Returns:
        dict: A dictionary containing the number of processed documents
              and a status message. Keys are 'processed' and 'message'.
              Returns an error message on failure.
    """
    client = None  # Initialize client to None for robust error handling
    embeddings_obj = None # Initialize embeddings to None
    num_processed = 0 # Keep track of processed docs count

    try:
        # 1. Initialize the Weaviate client using PipelineConfig
        logger.info("Attempting to get Weaviate client via PipelineConfig...")
        # Pass cfg if needed by your get_client implementation, otherwise remove it
        print(">>> ingest_block: About to call get_client()")
        client = PipelineConfig.get_client() # Call the corrected get_client

        # 2. --- ADDED CHECK ---
        # Explicitly check if get_client returned None (connection failure)
        if client is None:
            logger.critical("Failed to obtain a valid Weaviate client. Check connection details and Weaviate server status.")
            # Return an error dictionary suitable for the Flask route
            return {
                "processed": 0,
                "message": "Incremental ingestion failed: Could not connect to Weaviate."
            }
        # --- END ADDED CHECK ---

        # 3. Now it's safe to check is_live()
        logger.info("Weaviate client obtained. Checking if live...")
        if not client.is_live():
            logger.critical("Weaviate server is not responsive.")
            # Return an error dictionary
            return {
                "processed": 0,
                "message": "Incremental ingestion failed: Weaviate server is not responsive."
            }
        logger.info("Weaviate client is live.")

        # 4. Initialize embeddings via OllamaEmbeddings
        logger.info("Initializing embeddings...")
        try:
            embeddings_obj = OllamaEmbeddings(
                model=PipelineConfig.EMBEDDING_MODEL,
                base_url=PipelineConfig.EMBEDDING_BASE_URL
                # Add timeout or other params if needed
            )
            logger.info(f"Embeddings initialized with model: {PipelineConfig.EMBEDDING_MODEL}")
        except Exception as embed_e:
            logger.error(f"Error initializing embeddings: {embed_e}", exc_info=True)
            # Return an error dictionary
            return {
                "processed": 0,
                "message": f"Incremental ingestion failed: Could not initialize embeddings: {embed_e}"
            }

        # 5. Process new/modified documents incrementally
        logger.info(f"Starting incremental processing for directory: {folder}")
        incremental_processor = IncrementalDocumentProcessorBlock(data_dir=folder)
        new_docs = incremental_processor.process() # This loads, hashes, extracts, splits
        num_processed = len(new_docs) # Get count after processing

        # 6. Insert the resulting document chunks into Weaviate
        if new_docs:
            logger.info(f"Storing {num_processed} new document chunks into Weaviate...")
            store_documents(new_docs, client, embeddings_obj) # Ensure store_documents handles its own errors
        else:
            logger.info("No new or modified documents found to process.")

        # 7. Log success and prepare return value
        logger.info("Incremental ingestion completed: %d new document chunks processed.", num_processed)
        return {
            "processed": num_processed,
            "message": "Incremental ingestion completed successfully."
        }

    except Exception as e:
        # Catch any other unexpected errors during the process
        logger.critical(f"Error during incremental ingestion setup or execution: {e}", exc_info=True)
        # Return a generic error dictionary
        return {
            "processed": num_processed, # Report how many were processed before error if possible
            "message": f"Incremental ingestion failed: {str(e)}"
        }

    finally:
        # 8. Ensure the client connection is closed
        if client and hasattr(client, 'close') and callable(client.close):
            try:
                # Check if connected before closing, might prevent errors if already closed
                if client.is_connected():
                    client.close()
                    logger.info("Closed Weaviate client connection.")
                else:
                     logger.info("Weaviate client was already closed or not connected.")
            except Exception as close_err:
                logger.error(f"Error closing Weaviate client in finally block: {close_err}")
 

# --- Main Execution ---
def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--folder", "-f", type=str, default=cfg.paths.DOCUMENT_DIR,
                        help="Folder containing documents to ingest.")
    parser.add_argument("--verbose", "-v", action="store_true", help="Enable verbose logging.")
    args = parser.parse_args()
    if args.verbose:
        logger.setLevel(logging.DEBUG)
    # Initialize the Weaviate client using PipelineConfig
    try:
        client = PipelineConfig.get_client()
        if not client.is_live():
            raise Exception("Weaviate server is not responsive.")
    except Exception as e:
        logger.critical("Error connecting to Weaviate: %s", e)
        return
    # Initialize embeddings via OllamaEmbeddings
    try:
        from langchain_ollama import OllamaEmbeddings
        embeddings_obj = OllamaEmbeddings(
            model=PipelineConfig.EMBEDDING_MODEL,
            base_url=PipelineConfig.EMBEDDING_BASE_URL
        )
    except Exception as e:
        logger.error("Error initializing embeddings: %s", e)
        embeddings_obj = None
    # Process new/modified documents incrementally
    incremental_processor = IncrementalDocumentProcessorBlock(data_dir=args.folder)
    new_docs = incremental_processor.process()
    # Insert the resulting document chunks into Weaviate
    if embeddings_obj:
        store_documents(new_docs, client, embeddings_obj)
    else:
        logger.error("Embeddings object is not available; cannot compute embeddings.")
    logger.info("Incremental ingestion completed.")
    client.close()


# Only call main() if this script is run directly.
if __name__ == "__main__":
    main()

