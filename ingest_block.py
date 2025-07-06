# ingest_block.py (Cleaned and Optimized Version)

import os
import re
import csv
import json
import hashlib
import logging
import argparse
import traceback
import threading
import warnings
from pathlib import Path
from datetime import datetime, timezone
from typing import Any, Dict, Generator, List, Optional, Set, Tuple, Union
from collections import Counter
from concurrent.futures import ThreadPoolExecutor, as_completed
import weaviate
from weaviate.client import WeaviateClient



# Third-party libraries
import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.metrics.pairwise import cosine_similarity
from tenacity import retry, stop_after_attempt, wait_exponential

# Attempt to import optional dependencies
try:
    import pdfplumber
except ImportError:
    pdfplumber = None
    logging.warning("pdfplumber not installed; PDF table extraction might be limited.")

try:
    import docx
except ImportError:
    docx = None
    logging.warning("python-docx not installed; DOCX file processing will be unavailable.")

try:
    import pandas as pd
except ImportError:
    pd = None
    logging.warning("pandas not installed; table formatting might be basic.")

# LangChain & Ollama
from langchain_core.documents import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import OllamaEmbeddings  # Remove one if only one needed

# Weaviate v4
import weaviate
from weaviate.classes.data import DataObject

# Local modules
from config import cfg
from ingest_docs_v7 import PipelineConfig, RobustPDFLoaderV4
from centroid_manager import CentroidManager, should_recalculate_centroid

# Setup logging
logger = logging.getLogger(__name__)

# Initialize centroid manager
centroid_manager = CentroidManager(
    instance_alias=cfg.retrieval.WEAVIATE_ALIAS,
    collection_name=cfg.retrieval.COLLECTION_NAME,
    base_path=cfg.paths.CENTROID_DIR
)
_centroid_mgr = centroid_manager  # Optional alias if needed


def calculate_quality_score(
    text: str,
    tables: Optional[List[dict]] = None,
    cached_domain_keywords: Optional[Set[str]] = None
) -> float:
    """
    Unified document quality score:
      1. Info density
      2. Sentence-length penalty
      3. Keyword presence
      4. Structure (headings/paragraphs)
      5. Table content bonus
      6. Semantic relevance to domain centroid
    Returns value in [0.1, 1.0].
    """
    if not text or len(text) < 50:
        return 0.1

    clean = re.sub(r'\s+', ' ', text).strip()
    words = re.findall(r'\b\w+\b', clean.lower())
    if not words:
        return 0.1
    unique = set(words)
    info_density = len(unique) / len(words)

    sentences = [s.strip() for s in re.split(r'[.!?]+', clean) if s.strip()]
    if sentences:
        avg_len = sum(len(s.split()) for s in sentences) / len(sentences)
        sentence_score = min(1.0, max(0.0, 1.0 - abs(avg_len - 15) / 20))
    else:
        sentence_score = 0.2

    if cached_domain_keywords:
        matches = sum(1 for w in unique if w in cached_domain_keywords)
        keyword_score = min(1.0, matches / 10)
    else:
        keyword_score = 0.5

    has_head = bool(re.search(r'\n[A-Z][^.!?]*\n', text))
    has_para = '\n\n' in text
    structure_score = 0.3 + (0.4 if has_para else 0) + (0.3 if has_head else 0)

    table_score = min(1.0, len(tables) * 0.2) if tables else 0.0

    domain_rel = 0.0
    try:
        centroid = _centroid_mgr.get_centroid()
        if centroid is not None:
            emb = OllamaEmbeddings()
            vec = np.array(emb.embed_query(text)).reshape(1, -1)
            domain_rel = float(cosine_similarity(vec, [centroid])[0][0])
    except Exception:
        pass

    score = (
        info_density * 0.25 +
        sentence_score * 0.15 +
        keyword_score * 0.2 +
        structure_score * 0.1 +
        table_score * 0.1 +
        domain_rel * 0.2
    )
    return min(1.0, max(0.1, score))



def centroid_exists(centroid_path):
    return os.path.exists(centroid_path)


# Accepts an optional, already connected Weaviate client instance
# --- Corrected extract_text function ---

def extract_text(filepath: Path, client_instance: Optional[WeaviateClient] = None) -> str:
    """Extracts text from various file types, reusing a passed Weaviate client for PDFs."""
    ext = filepath.suffix.lower()
    logger.debug(f"Extracting text from: {filepath.name} (type: {ext})")

    if ext == ".txt":
        return extract_text_from_txt(filepath)
    if ext == ".docx":
        return extract_text_from_docx(filepath)
    if ext == ".csv":
        return extract_text_from_csv(filepath)
    if ext == ".md":
        return extract_text_from_txt(filepath)
    if ext == ".pdf":
        temp_client_created = False

        # Choose the client to use
        if client_instance and client_instance.is_live():
            loader_client = client_instance
            logger.debug(f"Reusing provided Weaviate client for PDF extraction: {filepath.name}")
        else:
            if client_instance:
                logger.warning(f"Provided Weaviate client not live, creating temp client for {filepath.name}")
            else:
                logger.warning(f"No Weaviate client passed, creating temp client for {filepath.name}")

            loader_client = PipelineConfig.get_client()
            if not loader_client or not loader_client.is_live():
                logger.error(f"Failed to obtain a live Weaviate client for PDF extraction: {filepath.name}")
                return ""
            temp_client_created = True
            logger.debug(f"Temporary Weaviate client created for PDF extraction: {filepath.name}")

        try:
            loader = RobustPDFLoaderV4(str(filepath), client=loader_client)
            docs = loader.load()
            return "\n".join(doc.page_content for doc in docs if doc.page_content)
        except Exception as e:
            logger.error(f"Error extracting PDF text {filepath.name}: {e}", exc_info=True)
            return ""
        finally:
            if temp_client_created:
                try:
                    loader_client.close()
                    logger.debug(f"Closed temporary Weaviate client after PDF extraction: {filepath.name}")
                except Exception as close_err:
                    logger.error(f"Error closing temporary Weaviate client for {filepath.name}: {close_err}")

    # Unsupported extension
    return ""


class CleanableOllamaEmbeddings(OllamaEmbeddings):
    def __del__(self):
        """Ensure client is closed when object is garbage collected"""
        self.close()
        
    def close(self):
        """Close the underlying HTTP client"""
        if hasattr(self, 'client') and self.client:
            self.client.close()
            logger.info("Closed Ollama embeddings client")

# --- Incremental Ingestion Class ---
class IncrementalDocumentProcessorBlock:
    def __init__(self, data_dir: str, meta_filename="ingested_files.json"):
        self.data_dir = Path(data_dir).resolve()
        if not self.data_dir.is_dir():
            logger.warning(f"Data directory {self.data_dir} does not exist. Creating.")
            try:
                self.data_dir.mkdir(parents=True, exist_ok=True)
            except OSError as e:
                 logger.error(f"Failed to create data directory {self.data_dir}: {e}")
                 raise # Stop if directory cannot be ensured
        # Initialize domain_keywords as empty set
        self.domain_keywords = set()
        self.meta_file = self.data_dir.parent / meta_filename 
        self.processed_files = self.load_metadata()

        
        self.domain_keywords = set(getattr(cfg.env, 'merged_keywords', []))
        logger.info(f"Cached {len(self.domain_keywords)} domain keywords for quality scoring")         
        self.meta_file = self.data_dir.parent / meta_filename # Store metafile outside data dir
        self.processed_files = self.load_metadata()
        self._load_domain_keywords()
    
    def _load_domain_keywords(self):
        """Load domain keywords from config, if available"""
        try:
            from config import cfg
            if hasattr(cfg, 'env') and hasattr(cfg.env, 'merged_keywords'):
                self.domain_keywords = set(cfg.env.merged_keywords)
                logger.info(f"Cached {len(self.domain_keywords)} domain keywords for quality scoring")
            else:
                logger.warning("No merged_keywords found in configuration")
        except (ImportError, AttributeError) as e:
            logger.warning(f"Could not load domain keywords: {e}")
            
    def load_metadata(self) -> dict:
        """Load a JSON file mapping file paths to their last processed hash."""
        if self.meta_file.exists():
            try:
                with open(self.meta_file, "r", encoding="utf-8") as f:
                    return json.load(f)
            except (json.JSONDecodeError, IOError) as e:
                logger.error(f"Error loading metadata from {self.meta_file}: {e}")
                # back up the corrupted file for inspection
                corrupt = self.meta_file.with_suffix(self.meta_file.suffix + ".corrupt")
                try:
                    os.replace(self.meta_file, corrupt)
                    logger.info(f"Backed up corrupted metadata to {corrupt}")
                except Exception as be:
                    logger.error(f"Failed to back up corrupt metadata: {be}")
                return {}
        logger.info(f"Metadata file {self.meta_file} not found. Starting fresh.")
        return {}

    def save_metadata(self) -> None:
       # write atomically to avoid half-corrupting the file
        tmp = self.meta_file.with_suffix(self.meta_file.suffix + ".tmp")
        try:
            with open(tmp, "w", encoding="utf-8") as f:
                json.dump(self.processed_files, f, indent=2)
            os.replace(tmp, self.meta_file)
        except IOError as e:
            logger.error(f"Error saving metadata to {self.meta_file}: {e}")
           # clean up leftover tmp
            try:
                if tmp.exists():
                    tmp.unlink()
            except Exception:
                pass

    def compute_hash(self, filepath: Path) -> str:
        """Compute the SHA-256 hash of a file’s contents."""
        hasher = hashlib.sha256()
        try:
            with open(filepath, "rb") as f:
                while True:
                    buf = f.read(65536) # Read in chunks
                    if not buf:
                        break
                    hasher.update(buf)
            return hasher.hexdigest()
        except IOError as e:
            logger.error(f"Error reading file for hash {filepath.name}: {e}")
            return "" # Return empty string on error

    def load_new_documents(self, client_instance: weaviate.Client) -> List[Document]:
        """Iterate over files, process new/updated ones with parallel processing."""
        docs = []
        counter_lock = threading.Lock()
        logger.info(f"Scanning directory: {self.data_dir}")
        
        # Collect files to process
        files_to_process = []
        for filepath in self.data_dir.glob("*"):
            if not filepath.is_file():
                continue
                    
            # Check file type support
            supported_extensions = set(f".{ext.lower()}" for ext in getattr(PipelineConfig, 'FILE_TYPES', ['pdf', 'txt', 'csv', 'docx', 'md']))
            if filepath.suffix.lower() not in supported_extensions:
                continue
                
            current_hash = self.compute_hash(filepath)
            if not current_hash:
                continue
                
            stored_hash = self.processed_files.get(str(filepath))
            if stored_hash == current_hash:
                continue
                
            files_to_process.append(filepath)
        
        logger.info(f"Found {len(files_to_process)} new/modified files to process")
        
        # Process files in parallel
        max_workers = min(os.cpu_count() or 4, 8)  # Limit to 8 threads max
        logger.info(f"Starting parallel processing with {max_workers} worker threads")
        processed_count = 0
        error_count = 0
        # Thread-safe counter
        
        
        def process_file(filepath):
            nonlocal processed_count, error_count
            logger.info(f"Processing file: {filepath.name}")
            text = extract_text(filepath, client_instance=client_instance)
            
            if text is None or not text.strip():
                with counter_lock:
                    self.processed_files[str(filepath)] = self.compute_hash(filepath)
                return None
                
            try:
                file_path_obj = Path(filepath)
                stat = file_path_obj.stat()
                created_dt = datetime.fromtimestamp(stat.st_ctime, tz=timezone.utc).isoformat().replace('+00:00', 'Z')
                modified_dt = datetime.fromtimestamp(stat.st_mtime, tz=timezone.utc).isoformat().replace('+00:00', 'Z')
                
                metadata = {
                    "source": str(filepath),
                    "filetype": file_path_obj.suffix.lower(),
                    "created_date": created_dt,
                    "modified_date": modified_dt,
                    "page": 0,
                    "content_flags": [],
                    "error_messages": [],
                    "quality_score": calculate_quality_score(text, self.domain_keywords)  # Pass cached keywords
                }
                
                doc = Document(page_content=text, metadata=metadata)
                
                with counter_lock:
                    self.processed_files[str(filepath)] = self.compute_hash(filepath)
                    processed_count += 1
                    
                return doc
            except Exception as e:
                with counter_lock:
                    error_count += 1
                logger.error(f"Error processing {filepath.name}: {e}")
                return None
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            future_to_file = {executor.submit(process_file, filepath): filepath for filepath in files_to_process}
            for future in as_completed(future_to_file):
                doc = future.result()
                if doc:
                    docs.append(doc)
    
        logger.info(f"Parallel processing complete. Used {max_workers} threads, Processed: {processed_count}, Errors: {error_count}")
        return docs

       
    def process(self, client_instance: weaviate.Client) -> List[Document]:
        """
        Load new documents, extract tables, calculate quality, filter,
        and then perform a centroid significance check & update.
        """
        docs: List[Document] = []

        # 1) Scan for new/modified files
        logger.info(f"Scanning directory: {self.data_dir}")
        files_to_process: List[Tuple[Path, str]] = []
        supported_exts = {
            f".{e.lower()}"
            for e in getattr(PipelineConfig, "FILE_TYPES", ["pdf", "txt", "csv", "docx", "md"])
        }
        for path in self.data_dir.glob("*"):
            if not path.is_file() or path.suffix.lower() not in supported_exts:
                continue
            h = self.compute_hash(path)
            if not h or self.processed_files.get(str(path)) == h:
                continue
            files_to_process.append((path, h))

        if not files_to_process:
            logger.info("No new/modified files to process.")
            return []

        # 2) Process files (parallel)
        batch_size  = self._calculate_optimal_batch_size([p for p, _ in files_to_process])
        max_workers = min(os.cpu_count() or 4, 8)
        logger.info(
            f"Processing {len(files_to_process)} files "
            f"in batches of {batch_size} with {max_workers} threads"
        )
        for batch in self._create_batches(files_to_process, batch_size):
            with ThreadPoolExecutor(max_workers=max_workers) as pool:
                futures = {
                    pool.submit(self._process_single_file, path, client_instance): h
                    for path, h in batch
                }
                for fut in as_completed(futures):
                    doc = fut.result()
                    h = futures[fut]
                    if doc:
                        docs.append(doc)
                        self.processed_files[doc.metadata["source"]] = h

        if not docs:
            logger.info("No new/modified documents yielded valid docs.")
            return []

        # 3) Split into chunks
        logger.info(f"Splitting {len(docs)} documents into chunks")
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=PipelineConfig.CHUNK_SIZE,
            chunk_overlap=PipelineConfig.CHUNK_OVERLAP
        )
        try:
            split_docs = splitter.split_documents(docs)
        except Exception as e:
            logger.error(f"Splitting error: {e}", exc_info=True)
            self.save_metadata()
            return []
        total_chunks = len(split_docs)
        logger.info(f"Generated {total_chunks} total chunks")

        # 4) Filter by length & quality (use live config values)
        min_len = PipelineConfig.MIN_CONTENT_LENGTH
        min_q   = cfg.ingestion.MIN_QUALITY_SCORE
        quality_docs = [
            d for d in split_docs
            if len(d.page_content) >= min_len
            and d.metadata.get("quality_score", 0) >= min_q
        ]
        rejected = total_chunks - len(quality_docs)
        if rejected:
            logger.info(
                f"Rejected {rejected} of {total_chunks} chunks "
                f"below quality threshold {min_q:.2f}"
            )

        self.save_metadata()
        if not quality_docs:
            return []

        # 5) Centroid significance check & update
        try:
            from langchain_community.embeddings import OllamaEmbeddings
            embedder = OllamaEmbeddings(
                model=PipelineConfig.EMBEDDING_MODEL,
                base_url=PipelineConfig.EMBEDDING_BASE_URL
            )
            new_vectors = [
                np.array(embedder.embed_query(d.page_content))
                for d in quality_docs
            ]
            if new_vectors:
                cm = CentroidManager(
                    instance_alias=cfg.retrieval.WEAVIATE_ALIAS,
                    collection_name=cfg.retrieval.COLLECTION_NAME,
                    base_path=cfg.paths.CENTROID_DIR
                )
                all_vectors  = cm.get_all_vectors(client_instance, cfg.retrieval.COLLECTION_NAME)
                old_centroid = cm.get_centroid()
                auto_thr     = cfg.ingestion.CENTROID_AUTO_THRESHOLD
                div_thr      = cfg.ingestion.CENTROID_DIVERSITY_THRESHOLD

                if should_recalculate_centroid(new_vectors, all_vectors, old_centroid, auto_thr, div_thr):
                    cm.update_centroid(all_vectors)
                    logger.info("Centroid recalculated and updated.")
                else:
                    logger.info("Centroid update skipped (insufficient change).")

        except Exception as e:
            logger.critical(f"Error during centroid update: {e}", exc_info=True)

        return quality_docs


                  

    def _process_single_file(self, filepath, client_instance):
        """Process a single file with enhanced extraction."""
        
        counter_lock = threading.Lock()
        logger.info(f"Processing file: {filepath.name}")
        
        # Extract text content
        text = extract_text(filepath, client_instance=client_instance)
        if text is None or not text.strip():
            with counter_lock:
                self.processed_files[str(filepath)] = self.compute_hash(filepath)
            return None
        
        # Extract tables
        tables = extract_tables(filepath)
        
        try:
            file_path_obj = Path(filepath)
            stat = file_path_obj.stat()
            created_dt = datetime.fromtimestamp(stat.st_ctime, tz=timezone.utc).isoformat().replace('+00:00', 'Z')
            modified_dt = datetime.fromtimestamp(stat.st_mtime, tz=timezone.utc).isoformat().replace('+00:00', 'Z')
            
            # Calculate quality score with tables and domain keywords
            quality_score = calculate_quality_score(
                text, 
                tables=tables,
                cached_domain_keywords=self.domain_keywords
            )
            
            metadata = {
                "source": str(filepath),
                "filetype": file_path_obj.suffix.lower(),
                "created_date": created_dt,
                "modified_date": modified_dt,
                "page": 0,
                "content_flags": [],
                "error_messages": [],
                "quality_score": quality_score,
                "has_tables": bool(tables),
                "table_count": len(tables) if tables else 0
            }
            
            # Add table data to metadata if present
            if tables:
                metadata["tables"] = tables
            
            doc = Document(page_content=text, metadata=metadata)
            
            with counter_lock:
                self.processed_files[str(filepath)] = self.compute_hash(filepath)
                
            return doc
            
        except Exception as e:
            with counter_lock:
                logger.error(f"Error processing {filepath.name}: {e}")
            return None

        
    # Enhance _calculate_optimal_batch_size in ingest_block.py:
    def _calculate_optimal_batch_size(self, files_to_process):
        """Calculate optimal batch size based on file types, sizes, and system resources"""
        if not files_to_process:
            return 10
            
        # Get system memory info
        try:
            import psutil
            available_memory = psutil.virtual_memory().available / (1024 * 1024 * 1024)  # GB
        except ImportError:
            available_memory = 8  # Default assumption: 8GB
            
        # Analyze file sizes and types
        total_size = 0
        file_types = Counter()
        
        for filepath in files_to_process:
            file_types[filepath.suffix.lower()] += 1
            try:
                total_size += filepath.stat().st_size
            except OSError:
                pass
                
        # Adjust batch size based on average file size and available memory
        avg_size_mb = (total_size / len(files_to_process)) / (1024 * 1024) if files_to_process else 0
        pdf_ratio = file_types.get('.pdf', 0) / len(files_to_process) if files_to_process else 0
        
        # Memory-based calculation
        if available_memory < 2:  # Less than 2GB available
            base_size = 2
        elif available_memory < 4:  # 2-4GB available
            base_size = 4
        else:  # More than 4GB available
            base_size = 8
            
        # Adjust for file characteristics
        if avg_size_mb > 10 or pdf_ratio > 0.7:  # Large files or mostly PDFs
            return max(2, base_size // 2)
        elif avg_size_mb > 1 or pdf_ratio > 0.3:  # Medium files or some PDFs
            return max(4, base_size)
        else:  # Small files, few PDFs
            return max(6, base_size * 2)

        
    def _create_batches(self, files, batch_size):
        """Create batches of files for processing."""
        return [files[i:i + batch_size] for i in range(0, len(files), batch_size)]

# --- Store Documents in Weaviate (Mostly Unchanged, Check Metadata) ---
def store_documents(docs: List[Document], client: weaviate.Client, embeddings_obj) -> None:
    """Stores documents in Weaviate using batch insertion."""
    if embeddings_obj is None:
        logger.error("store_documents: No embeddings object provided.")
        return

    if not docs:
        logger.info("store_documents: No documents to insert.")
        return

    collection_name = PipelineConfig.VECTORSTORE_COLLECTION
    try:
        collection = client.collections.get(collection_name)
    except Exception as e:
        logger.error(f"store_documents: Error fetching collection '{collection_name}': {e}", exc_info=True)
        return

    objects_to_insert: List[DataObject] = []
    logger.info(f"store_documents: Preparing {len(docs)} document chunks for insertion...")

    # --- Prepare objects with embeddings ---
    contents = [doc.page_content for doc in docs]
    embeddings = []
    if contents:
        try:
            # Embed in batches if embed_documents is efficient, otherwise loop embed_query
            if hasattr(embeddings_obj, 'embed_documents'):
                 logger.info(f"Embedding {len(contents)} chunks using embed_documents...")
                 embeddings = embeddings_obj.embed_documents(contents)
                 logger.info("Embeddings generated.")
            else:
                 logger.info(f"Embedding {len(contents)} chunks using embed_query loop...")
                 # Fallback to loop (potentially slower)
                 for i, content in enumerate(contents):
                     @retry(stop=stop_after_attempt(PipelineConfig.RETRY_ATTEMPTS),
                            wait=wait_exponential(multiplier=PipelineConfig.RETRY_WAIT_MULTIPLIER, max=PipelineConfig.RETRY_MAX_WAIT))
                     def get_single_embedding():
                         return embeddings_obj.embed_query(content)
                     embeddings.append(get_single_embedding())
                     if (i + 1) % 50 == 0: # Log progress
                          logger.info(f" Embedded {i+1}/{len(contents)} chunks...")
                 logger.info("Embeddings generated.")
        except Exception as embed_e:
             logger.error(f"Fatal error during embedding generation: {embed_e}", exc_info=True)
             return # Cannot proceed without embeddings

    if len(embeddings) != len(docs):
        logger.error(f"Mismatch between number of documents ({len(docs)}) and embeddings ({len(embeddings)}). Aborting insertion.")
        return

    # --- Create DataObject list ---
    for i, doc in enumerate(docs):
        # Ensure metadata keys required by Weaviate schema exist.
        # The keys were already prepared correctly in load_new_documents.
        properties = {
            "content": doc.page_content,
            "source": doc.metadata.get("source", "unknown"),
            "filetype": doc.metadata.get("filetype", "unknown"),
            "page": doc.metadata.get("page", 0), # Splitter might add page info
            "created_date": doc.metadata.get("created_date"), # Already ISO string
            "modified_date": doc.metadata.get("modified_date"), # Already ISO string
            "content_flags": doc.metadata.get("content_flags", []), # Use defaults if missing
            "error_messages": doc.metadata.get("error_messages", [])
        }
        # Filter out None values before insertion if schema doesn't allow nulls
        properties_filtered = {k: v for k, v in properties.items() if v is not None}

        objects_to_insert.append(
            DataObject(
                properties=properties_filtered,
                vector=embeddings[i] # Get corresponding embedding
            )
        )

    if not objects_to_insert:
        logger.warning("store_documents: No valid objects prepared for insertion after embedding.")
        return

    # --- Perform Batch Insertion ---
    logger.info(f"store_documents: Attempting to insert {len(objects_to_insert)} objects...")
    try:
        # Use a context manager for batching if available and preferred
        # Or use insert_many directly
        response = collection.data.insert_many(objects=objects_to_insert)

        logger.info("store_documents: insert_many operation finished.")
        # Process response (check for errors)
        if response and hasattr(response, 'has_errors') and response.has_errors:
            error_count = 0
            for index, error in response.errors.items():
                logger.error(f"  Index {index}: {error.message}")
                error_count += 1
            logger.error(f"store_documents: Batch insertion completed with {error_count} errors.")
        elif response:
            # Log success count if possible from response details
            success_count = len(objects_to_insert) - (len(response.errors) if hasattr(response, 'errors') else 0)
            logger.info(f"store_documents: Successfully inserted {success_count} objects.")
        else:
             logger.warning("store_documents: insert_many response was None or unexpected.")

    except Exception as e: # Catch Weaviate errors and others
        logger.error(f"store_documents: Batch insertion failed: {e}", exc_info=True)
        # Consider adding more specific error handling if needed


# --- OPTIMIZED run_ingestion (Root Fix, v2) ---
def run_ingestion(folder: str) -> dict:
    """
    Incrementally ingest files under `folder` into Weaviate,
    then perform a centroid significance check/update.

    Returns a JSON-serializable dict with:
      - new_vectors: List of lists (embeddings)
      - all_vectors: List of lists (all embeddings)
      - old_centroid: List (previous centroid)
      - message: Status string
    """
    # 1) Connect to Weaviate
    client = weaviate.connect_to_local(
        host=cfg.retrieval.WEAVIATE_HOST,
        port=cfg.retrieval.WEAVIATE_HTTP_PORT,
        grpc_port=cfg.retrieval.WEAVIATE_GRPC_PORT,
        skip_init_checks=True
    )

    try:
        # 2) Process new/modified files
        docs = IncrementalDocumentProcessorBlock(data_dir=folder).process(client)
        cm = CentroidManager()

        # Handle no new documents
        if not docs:
            logger.info("run_ingestion: No new or modified documents found.")
            # Explicitly handle None centroid
            old_cent = cm.get_centroid()
            if old_cent is None:
                old_cent = np.zeros(PipelineConfig.EMBED_DIM)
            return {
                "new_vectors": [],
                "all_vectors": [],
                "old_centroid": old_cent.tolist(),
                "message": "No new documents ingested."
            }

        # 3) Upsert & collect embeddings via the store_documents helper
        embedder = OllamaEmbeddings(
            model=PipelineConfig.EMBEDDING_MODEL,
            base_url=PipelineConfig.EMBEDDING_BASE_URL
        )
        # Insert everything in one shot
        store_documents(docs, client, embedder)

        # Meanwhile build new_vecs for centroid logic
        new_vecs = [
            np.array(embedder.embed_query(d.page_content))
            for d in docs
            ]
   

        # 4) Fetch all vectors & previous centroid
        all_vecs = cm.get_all_vectors(client, cfg.retrieval.COLLECTION_NAME)
        old_cent = cm.get_centroid()
        if old_cent is None:
            # default shape to first vec
            old_cent = np.zeros_like(all_vecs[0])

        # 5) Optional centroid update
        try:
            if should_recalculate_centroid(new_vecs, all_vecs, old_cent,
                                          cfg.ingestion.CENTROID_AUTO_THRESHOLD,
                                          cfg.ingestion.CENTROID_DIVERSITY_THRESHOLD):
                cm.update_centroid(all_vecs)
                logger.info("Centroid recalculated and updated.")
            else:
                logger.info("Centroid update skipped.")
        except Exception:
            logger.warning("Centroid update failed.", exc_info=True)

        # Prepare JSON-serializable output
        serial_new = [v.tolist() for v in new_vecs]
        serial_all = [v.tolist() for v in all_vecs]
        serial_old = old_cent.tolist()
        msg = f"Ingested {len(new_vecs)} vectors."
        logger.info(f"run_ingestion: {msg}")

        return {
            "new_vectors": serial_new,
            "all_vectors": serial_all,
            "old_centroid": serial_old,
            "message": msg
        }

    except Exception as e:
        logger.critical(f"run_ingestion error: {e}", exc_info=True)
        raise

    finally:
        client.close()




def after_ingestion(client, collection_name):
    """Update domain centroid after ingestion by averaging all document vectors."""
    logger.info(f"Updating domain centroid from collection '{collection_name}'...")
    
    # Use a lock to ensure thread safety
    with threading.Lock():
        try:
            vectors = []
            collection = client.collections.get(collection_name)
            
            # Log progress periodically
            processed = 0
            for obj in collection.iterator(include_vector=True, return_properties=[]):
                processed += 1
                if processed % 1000 == 0:
                    logger.info(f"Processed {processed} vectors...")
                
                if obj.vector and 'default' in obj.vector:
                    vectors.append(obj.vector['default'])
            
            if vectors:
                logger.info(f"Calculating centroid from {len(vectors)} vectors...")
                centroid = np.mean(np.array(vectors), axis=0)
                centroid_manager.save_centroid(centroid)
                logger.info(f"Domain centroid updated successfully with shape {centroid.shape}")
                return True
            else:
                logger.warning("No vectors found in collection. Centroid not updated.")
                return False
                
        except Exception as e:
            logger.error(f"Error updating centroid: {e}", exc_info=True)
            return False

def extract_text_from_txt(filepath: Path) -> str:
    try:
        import chardet
        with open(filepath, "rb") as f:
            raw_data = f.read()
            encoding = chardet.detect(raw_data)['encoding'] or 'utf-8'
        with open(filepath, "r", encoding=encoding) as f:
            content = f.read()
            logger.info(f"Successfully extracted {len(content)} characters from TXT file: {filepath.name}")
            return content
    except Exception as e:
        logger.error(f"Error reading TXT file {filepath.name}: {e}")
        return ""

# In ingest_block.py, enhance extract_text_from_docx:
def extract_text_from_docx(filepath: Path) -> str:
    if docx is None:
        logger.error(f"python-docx not installed. Cannot process DOCX: {filepath.name}")
        return ""
    
    try:
        doc = docx.Document(filepath)
        full_text = []
        
        # Extract main text with heading levels
        for para in doc.paragraphs:
            # Add heading level info if it's a heading
            if para.style.name.startswith('Heading'):
                level = para.style.name.replace('Heading ', '')
                full_text.append(f"{'#' * int(level)} {para.text}")
            else:
                full_text.append(para.text)
        
        # Extract tables with better formatting
        for table in doc.tables:
            table_text = []
            for i, row in enumerate(table.rows):
                cells = [cell.text for cell in row.cells]
                if i == 0:  # Header row
                    table_text.append("| " + " | ".join(cells) + " |")
                    table_text.append("| " + " | ".join(["---"] * len(cells)) + " |")
                else:
                    table_text.append("| " + " | ".join(cells) + " |")
            full_text.append("\n".join(table_text))
            
        return "\n\n".join(full_text)
    except Exception as e:
        logger.error(f"Error processing DOCX {filepath.name}: {e}")
        return ""


def extract_text_from_csv(filepath: Path) -> str:
    try:
        if pd is not None:
            df = pd.read_csv(filepath, encoding='utf-8', on_bad_lines='skip')
            # Convert to markdown table for better structure preservation
            return df.to_markdown(index=False)
        else:
            # Fallback to basic CSV reader
            with open(filepath, "r", encoding="utf-8", newline='') as f:
                reader = csv.reader(f)
                return "\n".join([", ".join(row) for row in reader])
    except Exception as e:
        logger.error(f"Error reading CSV file {filepath.name}: {e}")
        return ""

def extract_text_from_markdown(filepath: Path) -> str:
    try:
        with open(filepath, "r", encoding="utf-8") as f:
            content = f.read()
            # Preserve markdown structure
            return content
    except Exception as e:
        logger.error(f"Error reading Markdown file {filepath.name}: {e}")
        return ""
def extract_tables(filepath: Path) -> list:
    """Extract tables from various document formats."""
    ext = filepath.suffix.lower()
    tables = []
    
    try:
        if ext == '.pdf' and pdfplumber is not None:
            with pdfplumber.open(filepath) as pdf:
                for i, page in enumerate(pdf.pages):
                    for j, table in enumerate(page.extract_tables()):
                        if table:
                            tables.append({
                                "table_id": f"page_{i+1}_table_{j+1}",
                                "data": table,
                                "page": i+1
                            })
        elif ext == '.docx' and docx is not None:
            doc = docx.Document(filepath)
            for i, table in enumerate(doc.tables):
                data = []
                for row in table.rows:
                    data.append([cell.text for cell in row.cells])
                tables.append({
                    "table_id": f"table_{i+1}",
                    "data": data
                })
        elif ext == '.csv' and pd is not None:
            df = pd.read_csv(filepath)
            tables.append({
                "table_id": "csv_table",
                "data": df.values.tolist(),
                "headers": df.columns.tolist()
            })
        return tables
    except Exception as e:
        logger.error(f"Table extraction failed for {filepath.name}: {e}")
        return []
    
   
# --- Main Execution (for running script directly) ---
def main():
    import argparse
    import logging
    import sys

    def setup_logging(force_utf8=False, level=logging.INFO):
        stream_handler = logging.StreamHandler(sys.stdout)
        stream_handler.setFormatter(logging.Formatter('%(asctime)s - %(levelname)s - %(message)s'))
        if force_utf8:
            try:
                stream_handler.stream = open(sys.stdout.fileno(), mode='w', encoding='utf-8', buffering=1)
            except Exception:
                pass
        logging.basicConfig(handlers=[stream_handler], level=level, force=True)

    parser = argparse.ArgumentParser(description="Run incremental document ingestion.")
    parser.add_argument("--folder", "-f", type=str, default=cfg.paths.DOCUMENT_DIR if cfg else "./data",
                        help="Folder containing documents to ingest.")
    parser.add_argument("--verbose", "-v", action="store_true", help="Enable verbose logging.")
    args = parser.parse_args()

    log_level = logging.DEBUG if args.verbose else logging.INFO
    setup_logging(force_utf8=True, level=log_level)

    logger = logging.getLogger(__name__)
    logger.info("Running ingest_block script directly...")

    result = run_ingestion(args.folder)

    import pprint
    pprint.pprint(result)


if __name__ == "__main__":
    main()
